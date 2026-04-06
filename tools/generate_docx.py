"""Convert markdown docs to .docx using python-docx."""
import re
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── helpers ───────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "DDDDDD")
        tcBorders.append(border)
    tcPr.append(tcBorders)


def add_run_with_inline(para, text):
    """Add runs handling **bold** and `inline code` markers."""
    # split on **bold** and `code`
    pattern = re.compile(r"\*\*(.+?)\*\*|`(.+?)`")
    last = 0
    for m in pattern.finditer(text):
        # plain text before match
        if m.start() > last:
            r = para.add_run(text[last:m.start()])
            r.font.size = Pt(10.5)
        if m.group(1) is not None:  # **bold**
            r = para.add_run(m.group(1))
            r.bold = True
            r.font.size = Pt(10.5)
        else:  # `code`
            r = para.add_run(m.group(2))
            r.font.name = "Courier New"
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
        last = m.end()
    if last < len(text):
        r = para.add_run(text[last:])
        r.font.size = Pt(10.5)


def add_heading(doc, text, level):
    sizes = {1: 16, 2: 13, 3: 11}
    colors = {
        1: RGBColor(0xFF, 0x44, 0x1A),
        2: RGBColor(0x22, 0x22, 0x22),
        3: RGBColor(0x44, 0x44, 0x44),
    }
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(sizes.get(level, 11))
    run.font.color.rgb = colors.get(level, RGBColor(0x44, 0x44, 0x44))
    space_before = {1: 18, 2: 14, 3: 8}
    space_after  = {1: 6,  2: 4,  3: 2}
    p.paragraph_format.space_before = Pt(space_before.get(level, 8))
    p.paragraph_format.space_after  = Pt(space_after.get(level, 2))


def add_code_block(doc, lines):
    for line in lines:
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        p.paragraph_format.left_indent  = Cm(0.5)
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "F5F5F5")
        p._p.get_or_add_pPr().append(shd)


def add_table_md(doc, header, rows):
    col_n = len(header)
    table = doc.add_table(rows=1 + len(rows), cols=col_n)
    table.style = "Table Grid"
    # header
    for i, h in enumerate(header):
        cell = table.rows[0].cells[i]
        cell.text = h.strip()
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_bg(cell, "FF441A")
        set_cell_border(cell)
    # rows
    for r_i, row in enumerate(rows):
        for c_i, cell_text in enumerate(row):
            cell = table.rows[r_i + 1].cells[c_i]
            p = cell.paragraphs[0]
            add_run_with_inline(p, cell_text.strip())
            bg = "FFFFFF" if r_i % 2 == 0 else "FFF8F6"
            set_cell_bg(cell, bg)
            set_cell_border(cell)
    doc.add_paragraph()


def parse_table_line(line):
    return [c for c in line.strip().strip("|").split("|")]


def is_separator(line):
    return bool(re.match(r"^\|?[\s\-\|:]+\|?$", line.strip()))


# ── converter ─────────────────────────────────────────────────────────────────

def md_to_docx(md_path, docx_path, title):
    doc = Document()

    section = doc.sections[0]
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

    # Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_title.add_run(title)
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0xFF, 0x44, 0x1A)
    p_title.paragraph_format.space_after = Pt(4)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = p_sub.add_run("Rappi AI — Sistema de Análisis Inteligente de Operaciones")
    sub.font.size = Pt(11)
    sub.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    doc.add_paragraph()

    with open(md_path, encoding="utf-8") as f:
        lines = f.readlines()

    i = 0
    total = len(lines)

    while i < total:
        line = lines[i].rstrip("\n")

        # Skip first h1 (already shown as title)
        if line.startswith("# ") and i < 5:
            i += 1
            continue

        # Horizontal rule
        if re.match(r"^---+$", line.strip()):
            doc.add_paragraph()
            i += 1
            continue

        # Headings
        if line.startswith("### "):
            add_heading(doc, line[4:].strip(), 3)
            i += 1
            continue
        if line.startswith("## "):
            add_heading(doc, line[3:].strip(), 2)
            i += 1
            continue
        if line.startswith("# "):
            add_heading(doc, line[2:].strip(), 1)
            i += 1
            continue

        # Code block
        if line.startswith("```"):
            code_lines = []
            i += 1
            while i < total and not lines[i].startswith("```"):
                code_lines.append(lines[i].rstrip("\n"))
                i += 1
            add_code_block(doc, code_lines)
            i += 1
            continue

        # Blockquote
        if line.startswith("> "):
            p = doc.add_paragraph()
            run = p.add_run(line[2:].strip())
            run.italic = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)
            p.paragraph_format.left_indent = Cm(0.8)
            i += 1
            continue

        # Table
        if line.startswith("|"):
            header = parse_table_line(line)
            i += 1
            if i < total and is_separator(lines[i]):
                i += 1
            rows = []
            while i < total and lines[i].strip().startswith("|"):
                rows.append(parse_table_line(lines[i]))
                i += 1
            add_table_md(doc, header, rows)
            continue

        # Numbered list
        if re.match(r"^\d+\. ", line):
            p = doc.add_paragraph(style="List Number")
            add_run_with_inline(p, re.sub(r"^\d+\. ", "", line))
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(1)
            i += 1
            continue

        # Bullet list
        if line.startswith("- ") or line.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            add_run_with_inline(p, line[2:])
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(1)
            i += 1
            continue

        # Empty line
        if not line.strip():
            i += 1
            continue

        # Normal paragraph
        p = doc.add_paragraph()
        add_run_with_inline(p, line)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(4)
        i += 1

    doc.save(docx_path)
    print(f"Saved: {docx_path}")


# ── run ───────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    import os
    base = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))

    md_to_docx(
        os.path.join(base, "docs", "documentacion_tecnica.md"),
        os.path.join(base, "docs", "documentacion_tecnica.docx"),
        "Documentación Técnica",
    )

    md_to_docx(
        os.path.join(base, "README.md"),
        os.path.join(base, "docs", "README.docx"),
        "README — Guía Rápida del Proyecto",
    )

    print("Done.")
